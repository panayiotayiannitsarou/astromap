import io
from pathlib import Path

from docx import Document


REFERENCE_DIR = Path(__file__).resolve().parent / "references"
DEFAULT_INSTRUCTIONS = REFERENCE_DIR / "Odigies_v4.docx"
DEFAULT_STYLE = REFERENCE_DIR / "Elena_style_example.docx"


def docx_text(source) -> str:
    """Extract paragraphs and tables from a DOCX path or uploaded bytes."""
    if isinstance(source, (bytes, bytearray)):
        source = io.BytesIO(source)
    document = Document(source)
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            line = " | ".join(cell.text.strip() for cell in row.cells)
            if line.strip(" |"):
                blocks.append(line)
    return "\n".join(blocks)


def load_default_references() -> tuple[str, str]:
    if not DEFAULT_INSTRUCTIONS.exists() or not DEFAULT_STYLE.exists():
        raise FileNotFoundError("Λείπουν οι ενσωματωμένες οδηγίες v4 ή το πρότυπο ύφους.")
    return docx_text(DEFAULT_INSTRUCTIONS), docx_text(DEFAULT_STYLE)
