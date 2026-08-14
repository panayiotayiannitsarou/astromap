from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from prompts import fmt
from astrology import movement_text

def _shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def build_audit_docx(chart, personal, prompt):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
    styles=d.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(10.5)
    for s,size,color in [('Title',26,'1D3A34'),('Heading 1',18,'1D3A34'),('Heading 2',14,'5B7F6A')]:
        styles[s].font.name='Aptos Display'; styles[s].font.size=Pt(size); styles[s].font.color.rgb=RGBColor.from_string(color)
    title=d.add_paragraph(style='Title'); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.add_run('AstroCheck — Δελτίο ελέγχου')
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(chart.name).bold=True
    d.add_heading('Βασικά δεδομένα',1)
    table=d.add_table(rows=0,cols=2); table.style='Table Grid'
    for a,b in [('Ημερομηνία',chart.date),('Ώρα',chart.time),('Τόπος',chart.place),('Σύστημα Οίκων',chart.house_system)]:
        cells=table.add_row().cells; cells[0].text=a; cells[1].text=b; _shade(cells[0],'E5EFE8')
    d.add_heading('Πλανήτες και σημεία',1)
    t=d.add_table(rows=1,cols=4); t.style='Table Grid'
    for i,h in enumerate(['Σημείο','Θέση','Οίκος','Κίνηση']): t.rows[0].cells[i].text=h; _shade(t.rows[0].cells[i],'1D3A34'); t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255)
    for pnt in chart.points:
        c=t.add_row().cells; c[0].text=pnt.name; c[1].text=fmt(pnt); c[2].text=str(pnt.house or '—'); c[3].text=movement_text(pnt)
    d.add_heading('Υποχρεωτικός έλεγχος τετραγώνων και αντιθέσεων',1)
    hard=[a for a in chart.aspects if a.aspect in ('Τετράγωνο','Αντίθεση')]
    t=d.add_table(rows=1,cols=4); t.style='Table Grid'
    for i,h in enumerate(['Ζεύγος','Όψη','Orb','Βαρύτητα']): t.rows[0].cells[i].text=h; _shade(t.rows[0].cells[i],'1D3A34'); t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255)
    for a in hard:
        c=t.add_row().cells; c[0].text=f'{a.first}–{a.second}'; c[1].text=a.aspect; c[2].text=a.orb_text; c[3].text=a.weight
    d.add_page_break(); d.add_heading('Πλήρης εντολή για δημιουργία ανάλυσης',1); d.add_paragraph(prompt)
    bio=BytesIO(); d.save(bio); return bio.getvalue()

def build_analysis_docx(title_name, analysis):
    d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
    d.styles['Normal'].font.name='Aptos'; d.styles['Normal'].font.size=Pt(10.5)
    for s,size,color in [('Title',26,'1D3A34'),('Heading 1',18,'1D3A34'),('Heading 2',14,'5B7F6A'),('Heading 3',12,'5B7F6A')]:
        d.styles[s].font.name='Aptos Display'; d.styles[s].font.size=Pt(size); d.styles[s].font.color.rgb=RGBColor.from_string(color)
    p=d.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Πλήρης Αστρολογική Ανάλυση')
    s=d.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER; s.add_run(title_name).bold=True
    for raw in analysis.splitlines():
        line=raw.strip()
        if not line: d.add_paragraph(); continue
        if line.startswith('### '): d.add_heading(line[4:],3)
        elif line.startswith('## '): d.add_heading(line[3:],2)
        elif line.startswith('# '): d.add_heading(line[2:],1)
        elif line.startswith(('- ','• ')): d.add_paragraph(line[2:],style='List Bullet')
        elif line[:3].rstrip('.').isdigit() and '. ' in line[:5]: d.add_paragraph(line.split('. ',1)[1],style='List Number')
        else: d.add_paragraph(line)
    bio=BytesIO(); d.save(bio); return bio.getvalue()
